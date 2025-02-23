import docx
import asyncio
from googletrans import Translator
from serbian_text_converter import SerbianTextConverter


async def translate_paragraph(paragraph, translator, target_lang='en'):
    print(f"Processing paragraph: {paragraph.text}")  # Log the paragraph text being processed
    translated_runs = []
    current_run_text = ''  # This will hold the accumulated text from multiple runs

    # Iterate through all the runs in the paragraph
    for run in paragraph.runs:
        if run.text.strip():  # If there's any text to translate
            print(f"Accumulating text: {run.text}")  # Debugging output to see the accumulated text
            current_run_text += run.text  # Collect the text from the current run
        else:
            # Add the non-translated run (e.g., spaces or line breaks)
            translated_runs.append(run)

    # If any text was accumulated, translate it as a whole
    if current_run_text:
        try:
            # Await the translation coroutine
            translated_text = await translator.translate(current_run_text, dest=target_lang)
            print(f"Translated text: {translated_text.text}")

            # Create a new run for the translated text and preserve the original formatting
            translated_run = paragraph.add_run(translated_text.text)

            # Preserve formatting (bold, italic, underline, font size, color, etc.)
            translated_run.bold = run.bold
            translated_run.italic = run.italic
            translated_run.underline = run.underline
            translated_run.font.size = run.font.size
            translated_run.font.color.rgb = run.font.color.rgb
            translated_run.font.name = run.font.name
            translated_run.font.highlight_color = run.font.highlight_color

            # If translating to Serbian Latin, convert Cyrillic to Latin
            if target_lang == "sr_Latn":
                translated_run.text = SerbianTextConverter.to_latin(translated_run.text)

            # Ensure space is added if necessary after translated text
            if translated_run.text and not translated_run.text.endswith(' '):
                translated_run.text += ' '

            translated_runs.append(translated_run)
        except Exception as e:
            print(f"Error translating text: {current_run_text}, Error: {e}")

    # Clear the paragraph and rebuild it with translated runs
    paragraph.clear()

    # Re-add the translated runs to the paragraph
    for translated_run in translated_runs:
        paragraph.add_run(translated_run.text).bold = translated_run.bold
        paragraph.runs[-1].italic = translated_run.italic
        paragraph.runs[-1].underline = translated_run.underline
        paragraph.runs[-1].font.size = translated_run.font.size
        paragraph.runs[-1].font.color.rgb = translated_run.font.color.rgb
        paragraph.runs[-1].font.name = translated_run.font.name
        paragraph.runs[-1].font.highlight_color = translated_run.font.highlight_color

    return paragraph


# Function to process all paragraphs (including tables, headers, and footers)
async def process_paragraphs(doc, translator, target_lang='en'):
    # Process all paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Only process non-empty paragraphs
            await translate_paragraph(paragraph, translator, target_lang)
        else:
            print(f"Skipped empty paragraph.")  # Log if we skip an empty paragraph

    # Process tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Translate the text in each cell
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():  # Only process non-empty paragraphs in cells
                        await translate_paragraph(paragraph, translator, target_lang)
                    else:
                        print(f"Skipped empty cell paragraph.")  # Log skipped empty cell paragraph

    # Process headers and footers (if any)
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    await translate_paragraph(paragraph, translator, target_lang)
                else:
                    print(f"Skipped empty header paragraph.")  # Log skipped empty header paragraph
        if section.footer:
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip():
                    await translate_paragraph(paragraph, translator, target_lang)
                else:
                    print(f"Skipped empty footer paragraph.")  # Log skipped empty footer paragraph


# Function to translate the entire document
async def translate_document(input_file, output_file, target_lang='en'):
    # Load the Word document
    doc = docx.Document(input_file)
    translator = Translator()

    # Process paragraphs, tables, headers, and footers
    await process_paragraphs(doc, translator, target_lang)

    # Save the translated document
    doc.save(output_file)


# Create a function to run the translation process without closing the loop prematurely
def run_translation(input_path, output_path, target_lang):
    try:
        # Check if there is already a running event loop (this try loop is only needed because of PyCharm)
        try:
            loop = asyncio.get_running_loop()  # Try to get the current running loop
        except RuntimeError:  # If no loop is running, create a new one
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

        # Run the translation function
        loop.run_until_complete(translate_document(input_path, output_path, target_lang))
        print(f"Document has been translated and saved as {output_path}.")
    except Exception as e:
        print(f"Error occurred: {e}")


# Run the translation process
if __name__ == '__main__':
    input_path = 'dokument.docx'  # Path to your Word document
    output_path = 'translated_output.docx'  # Path to save the translated document
    target_lang = 'sr_Latn'  # Target language for translation | sr - Српски | sr_Latn - Srpski

    # Running the translation process manually
    run_translation(input_path, output_path, target_lang)
