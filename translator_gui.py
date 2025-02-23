import tkinter as tk
from tkinter import filedialog, messagebox
import docx
import asyncio
from googletrans import Translator
from serbian_text_converter import SerbianTextConverter

# Google Translate language dictionary
LANGUAGES = {
    "en": "English",
    "sr": "Serbian (Cyrillic)",
    "sr_Latn": "Serbian (Latin)",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
    "it": "Italian",
    "ru": "Russian",
    "pt": "Portuguese",
    "zh-cn": "Chinese (Simplified)",
    "ja": "Japanese",
    "ko": "Korean",
    "ar": "Arabic",
    # Add more languages as necessary
}

class TranslationApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Translator")
        self.root.geometry("700x300")
        self.root.config(bg="#2e2e2e")  # Dark background

        # File path variables
        self.input_path = tk.StringVar()
        self.output_path = tk.StringVar()
        self.output_name = tk.StringVar()

        # Target language variables
        self.target_language_code = tk.StringVar()  # Store language code instead of name
        self.target_language_name = tk.StringVar()  # Store the language name for display
        self.target_language_code.set("en")  # Default to English code (en)
        self.target_language_name.set(LANGUAGES["en"])  # Default to "English" as the name

        # Title label
        self.title_label = tk.Label(self.root, text="Word Document Translator", font=("Arial", 16), fg="white", bg="#2e2e2e")
        self.title_label.pack(pady=10)

        # File selection frame
        self.file_frame = tk.Frame(self.root, bg="#2e2e2e")
        self.file_frame.pack(pady=20)

        self.input_file_label = tk.Label(self.file_frame, text="Select Word file:", font=("Arial", 12), fg="white", bg="#2e2e2e")
        self.input_file_label.pack(side=tk.LEFT, padx=5)

        self.input_file_button = tk.Button(self.file_frame, text="Browse", command=self.select_file, bg="#4a4a4a", fg="white")
        self.input_file_button.pack(side=tk.LEFT)

        self.input_file_entry = tk.Entry(self.file_frame, textvariable=self.input_path, width=40, font=("Arial", 12))
        self.input_file_entry.pack(side=tk.LEFT, padx=5)

        # Language selection frame
        self.language_frame = tk.Frame(self.root, bg="#2e2e2e")
        self.language_frame.pack(pady=10)

        self.language_label = tk.Label(self.language_frame, text="Select target language:", font=("Arial", 12),
                                       fg="white", bg="#2e2e2e")
        self.language_label.pack(side=tk.LEFT, padx=5)

        # Create a list of language names (for display) and their corresponding language codes (for values)
        language_codes = list(LANGUAGES.keys())  # Language codes for values
        language_names = list(LANGUAGES.values())  # Language names for display

        # Function to update language code based on the selected language name
        def update_language_code(*args):
            selected_language_name = self.target_language_name.get()
            for code, name in LANGUAGES.items():
                if name == selected_language_name:
                    self.target_language_code.set(code)
                    break

        # Trace the target_language_name variable so that update_language_code gets called when the selection changes
        self.target_language_name.trace("w", update_language_code)

        # Create the OptionMenu using the language names (displayed) and corresponding language codes (values)
        self.language_dropdown = tk.OptionMenu(self.language_frame, self.target_language_name, *language_names)
        self.language_dropdown.config(font=("Arial", 12), bg="#4a4a4a", fg="white")
        self.language_dropdown.pack(side=tk.LEFT)

        # Output path and filename
        self.output_frame = tk.Frame(self.root, bg="#2e2e2e")
        self.output_frame.pack(pady=10)

        self.output_label = tk.Label(self.output_frame, text="Select output location and name:", font=("Arial", 12), fg="white", bg="#2e2e2e")
        self.output_label.pack(side=tk.LEFT, padx=5)

        self.output_name_entry = tk.Entry(self.output_frame, textvariable=self.output_name, width=40, font=("Arial", 12))
        self.output_name_entry.pack(side=tk.LEFT, padx=5)

        self.output_path_button = tk.Button(self.output_frame, text="Browse", command=self.select_output_location, bg="#4a4a4a", fg="white")
        self.output_path_button.pack(side=tk.LEFT)

        # Translate button
        self.translate_button = tk.Button(self.root, text="Translate", command=self.translate_document, bg="#4a4a4a", fg="white", font=("Arial", 14))
        self.translate_button.pack(pady=20)

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if file_path:
            self.input_path.set(file_path)

            # Automatically update the output path
            output_filename = file_path.split("/")[-1].replace(".docx", "") + "-translated.docx"
            output_dir = "/".join(file_path.split("/")[:-1])  # Get the directory of the selected file
            output_path = f"{output_dir}/{output_filename}"
            self.output_path.set(output_path)  # Set the output path

            # Ensure the output text box displays the output path correctly
            self.output_name.set(f"{output_dir}/{output_filename}")  # You can also update the filename part if needed

    def select_output_location(self):
        output_dir = filedialog.askdirectory()
        if output_dir:
            # Default to the same directory as input file if not specified
            output_filename = self.input_path.get().split("/")[-1].replace(".docx", "") + self.output_name.get()
            self.output_path.set(f"{output_dir}/{output_filename}")

    async def translate_paragraph(self, paragraph, translator, target_lang='en'):
        print(f"Processing paragraph: {paragraph.text}")  # Log the paragraph text being processed
        translated_runs = []
        current_run_text = ''  # This will hold the accumulated text from multiple runs

        # Iterate through all the runs in the paragraph
        for run in paragraph.runs:
            if run.text.strip():  # If there's any text to translate
                print(f"Accumulating text: {run.text}")  # Debugging output to see the accumulated text
                current_run_text += run.text  # Collect the text from the current run
            else:
                # Add the non-translated run (e.g., spaces or line breaks)
                translated_runs.append(run)

        # If any text was accumulated, translate it as a whole
        if current_run_text:
            try:
                # Await the translation coroutine
                translated_text = await translator.translate(current_run_text, dest=target_lang)
                print(f"Translated text: {translated_text.text}")

                # Create a new run for the translated text and preserve the original formatting
                translated_run = paragraph.add_run(translated_text.text)

                # Preserve formatting (bold, italic, underline, font size, color, etc.)
                translated_run.bold = run.bold
                translated_run.italic = run.italic
                translated_run.underline = run.underline
                translated_run.font.size = run.font.size
                translated_run.font.color.rgb = run.font.color.rgb
                translated_run.font.name = run.font.name
                translated_run.font.highlight_color = run.font.highlight_color

                # If translating to Serbian Latin, convert Cyrillic to Latin
                if target_lang == "sr_Latn":
                    translated_run.text = SerbianTextConverter.to_latin(translated_run.text)

                # Ensure space is added if necessary after translated text
                if translated_run.text and not translated_run.text.endswith(' '):
                    translated_run.text += ' '

                translated_runs.append(translated_run)
            except Exception as e:
                print(f"Error translating text: {current_run_text}, Error: {e}")

        # Clear the paragraph and rebuild it with translated runs
        paragraph.clear()

        # Re-add the translated runs to the paragraph
        for translated_run in translated_runs:
            paragraph.add_run(translated_run.text).bold = translated_run.bold
            paragraph.runs[-1].italic = translated_run.italic
            paragraph.runs[-1].underline = translated_run.underline
            paragraph.runs[-1].font.size = translated_run.font.size
            paragraph.runs[-1].font.color.rgb = translated_run.font.color.rgb
            paragraph.runs[-1].font.name = translated_run.font.name
            paragraph.runs[-1].font.highlight_color = translated_run.font.highlight_color

        return paragraph

    async def process_paragraphs(self, doc, translator, target_lang='en'):
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only process non-empty paragraphs
                await self.translate_paragraph(paragraph, translator, target_lang)

        # Process tables in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Translate the text in each cell
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():  # Only process non-empty paragraphs in cells
                            await self.translate_paragraph(paragraph, translator, target_lang)

        # Process headers and footers (if any)
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        await self.translate_paragraph(paragraph, translator, target_lang)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        await self.translate_paragraph(paragraph, translator, target_lang)

    def translate_document(self):
        input_path = self.input_path.get()
        output_path = self.output_path.get()
        target_lang = self.target_language_code.get()

        if not input_path or not output_path or not target_lang:
            messagebox.showerror("Error", "Please fill out all fields.")
            return

        try:
            # We need to run the async function inside asyncio
            asyncio.run(self.translate_document_async(input_path, output_path, target_lang))
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")

    async def translate_document_async(self, input_path, output_path, target_lang):
        try:
            # Load the Word document
            doc = docx.Document(input_path)
            translator = Translator()

            # Process paragraphs, tables, headers, and footers
            await self.process_paragraphs(doc, translator, target_lang)

            # Save the translated document
            doc.save(output_path)
            messagebox.showinfo("Success", f"Document has been translated and saved as {output_path}.")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")

# Create and run the GUI
if __name__ == '__main__':
    root = tk.Tk()
    app = TranslationApp(root)
    root.mainloop()
