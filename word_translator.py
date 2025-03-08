import threading
import docx
from azure.ai.translation.text import TextTranslationClient
from azure.core.credentials import AzureKeyCredential


class WordTranslationApp:
    def __init__(self, input_path, output_path, target_lang="en", progress_callback=None):
        self.input_path = input_path
        self.output_path = output_path
        self.target_language_code = target_lang
        self.progress_callback = progress_callback

        # Set up Azure Translator credentials
        self.endpoint = self.endpoint = "https://api.cognitive.microsofttranslator.com/"
        self.subscription_key = "Cbp6CjRkzbt1WdI5taT02TFvCUms0omfSKUIKJ5O6aUaIw3dprGCJQQJ99BCAC5RqLJXJ3w3AAAbACOGZYQb"
        self.region = "westeurope"
        self.client = TextTranslationClient(
            endpoint=self.endpoint,
            credential=AzureKeyCredential(self.subscription_key),
            headers={"Ocp-Apim-Subscription-Region": self.region}
        )

        # Start the translation process
        self.translate_document()

    def translate_paragraph(self, paragraph, client, target_lang='en'):
        """Translate a single paragraph."""
        print(f"Processing paragraph: {paragraph.text}")
        translated_runs = []
        current_run_text = ''

        for run in paragraph.runs:
            if run.text.strip():
                current_run_text += run.text
            else:
                translated_runs.append(run)

        if current_run_text:
            try:
                response = client.translate(
                    body=[current_run_text], to_language=[target_lang]
                )
                translated_text = response[0].translations[0].text
                print(f"Translated text: {translated_text}")

                translated_run = paragraph.add_run(translated_text)
                # Preserve formatting from the last run
                if paragraph.runs:
                    last_run = paragraph.runs[-1]
                    translated_run.bold = last_run.bold
                    translated_run.italic = last_run.italic
                    translated_run.underline = last_run.underline
                    translated_run.font.size = last_run.font.size
                    translated_run.font.color.rgb = last_run.font.color.rgb
                    translated_run.font.name = last_run.font.name
                    translated_run.font.highlight_color = last_run.font.highlight_color

                translated_runs.append(translated_run)
            except Exception as e:
                print(f"Error translating text: {current_run_text}, Error: {e}")

        # Clear and rebuild the paragraph
        paragraph.clear()
        for translated_run in translated_runs:
            new_run = paragraph.add_run(translated_run.text)
            new_run.bold = translated_run.bold
            new_run.italic = translated_run.italic
            new_run.underline = translated_run.underline
            new_run.font.size = translated_run.font.size
            new_run.font.color.rgb = translated_run.font.color.rgb
            new_run.font.name = translated_run.font.name
            new_run.font.highlight_color = translated_run.font.highlight_color

        return paragraph

    def process_paragraphs(self, doc, client, target_lang='en'):
        threads = []
        total_paragraphs = len(doc.paragraphs)
        progress_counter = [0]

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                thread = threading.Thread(target=self.translate_paragraph, args=(paragraph, client, target_lang))
                thread.start()
                threads.append((thread, paragraph))

        for thread, paragraph in threads:
            thread.join()
            progress_counter[0] += 1
            if self.progress_callback:
                progress = int((progress_counter[0] / total_paragraphs) * 100)
                self.progress_callback(progress)

    def translate_document(self):
        if not self.input_path or not self.output_path or not self.target_language_code:
            print("Error: Please provide all required paths and target language.")
            return

        try:
            doc = docx.Document(self.input_path)
            self.process_paragraphs(doc, self.client, self.target_language_code)
            doc.save(self.output_path)
            print(f"Document has been translated and saved as {self.output_path}.")
        except Exception as e:
            print(f"Error: An error occurred: {e}")
            if self.progress_callback:
                self.progress_callback(0)
